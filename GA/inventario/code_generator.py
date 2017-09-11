from GA import settings

from docx import Document
from docx.shared import Pt


def generate_full_code(user, product, number):
    serie = 0
    if number >= 100000000:
        serie = number / 100000000
        number = number % 100000000
    enterprise = settings.ENTERPRISE_DIC[user.enterprise]
    country = settings.COUNTRY_DIC[user.country]
    city = settings.CITY_DIC[user.city]
    dept = settings.DEPARTMENTS_DIC[product.department]
    code = product.code

    return (
        enterprise +
        country + '-' +
        city + '-' +
        dept + '-' +
        code + '-' +
        'S' + format(serie, '05d') + '-' +
        format(number, '08d')
    )

def create_codes_file(product_class, products, start, end, code_range):
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(8)
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = 'Code Row 1'
    table.rows[0].cells[1].text = 'Code Row 2'
    cell_index = 0
    counter = 0
    cell = table.add_row().cells
    for product in products:
        counter += 1
        if cell_index <= 1:
            if product.number == start or product.number == end:
                cell[cell_index].text = product.full_code
                cell_index += 1

            elif (counter % code_range) == 0:
                cell[cell_index].text = product.full_code
                cell_index += 1
        else:
            cell_index = 0
            cell = table.add_row().cells
            if product.number == start or product.number == end:
                cell[cell_index].text = product.full_code
                cell_index += 1

            elif (counter % code_range) == 0:
                cell[cell_index].text = product.full_code
                cell_index += 1
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.style = document.styles['Normal']
    return document